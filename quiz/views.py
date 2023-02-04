from random import sample
from . import forms
from django.db.models import Q
import openpyxl
import os
from django.utils import timezone
from django.http import HttpResponse, Http404
from django.conf import settings
from django.shortcuts import render, redirect
from accounts.models import Results, Temporary_user
from . import models
from docx import Document
from docx.shared import Inches, Pt, RGBColor

from datetime import datetime

# Create your views here.


def to_minut(nums) -> str:
    if int(nums) <= 60:
        return f"{nums} sekund"
    else:
        return f"{nums//60} minut {nums%60} sekund"


def ExamsView(request):
    exams = models.Exams.objects.filter(status="open")
    context = {'exams': exams}
    return render(request, 'home.html', context)


t = None


def ExamsDetailView(request, pk):
    print('#'*10)
    print(request)
    print('#'*10)
    context = {}
    choose_exam = models.Exams.objects.get(pk=pk)
    exam_quizes = list(models.Quiz.objects.filter(exam=pk))
    choose_exam_quizes = sample(exam_quizes, choose_exam.questions_count)
    user = request.user

    if request.method == "GET":
        print("GET")
        now = timezone.now()
        start_time = datetime.strftime(now, '%Y-%m-%d %H:%M:%S')
        request.session["start_time"] = start_time
        context = {
            'total': [i for i in range(1, len(list(choose_exam_quizes))+1)],
            'choose_exam': choose_exam,
            'choose_exam_quizes': choose_exam_quizes,
            'time': choose_exam.get_all_time() * 60,

        }
        return render(request, 'exams_items.html', context)

    elif request.method == "POST":
        end = timezone.now()
        end_time_str = datetime.strftime(end, '%Y-%m-%d %H:%M:%S')
        request.session["end_time"] = end_time_str
        print(request.session.get('end_time'), 'songi vaqt')

        start_time_str = request.session.get("start_time")
        start_time = datetime.strptime(start_time_str, '%Y-%m-%d %H:%M:%S')
        end_time = datetime.strptime(end_time_str, '%Y-%m-%d %H:%M:%S')
        duration = end_time - start_time
        total_seconds = to_minut(duration.seconds)
        request.session['total_seconds'] = total_seconds

        questions = choose_exam_quizes
        wrong, correct, total = 0, 0, 0

        file = Document()
        style = file.styles['Normal']
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(14)
        font.color.rgb = RGBColor(51, 0, 0)

        # TEMPRARY USER
        temprary_first_name = request.POST.get('temprary_first_name')
        temprary_last_name = request.POST.get('temprary_last_name')

        i = 1
        for question in questions:
            total += 1
            var_a = question.answer_a
            var_b = question.answer_b
            var_c = question.answer_c
            var_d = question.answer_d

            file.add_heading(f"{i}. {question.question}", level=1)
            i += 1
            choices = {
                'answer_a': ["a) ", var_a],
                'answer_b': ["b) ", var_b],
                'answer_c': ["c) ", var_c],
                'answer_d': ["d) ", var_d]
            }
            for choice, write in choices.items():
                if question.answer == choice:
                    para = file.add_paragraph(style="List 2").add_run(
                        f"{write[0]+write[1]}")
                    para.font.color.rgb = RGBColor(0, 255, 0)

                elif request.POST.get(question.id) == choice:
                    para = file.add_paragraph(style="List 2").add_run(
                        f"{write[0]+write[1]}")
                    para.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    para = file.add_paragraph(style="List 2").add_run(
                        f"{write[0]+write[1]}")
                    para.font.color.rgb = RGBColor(0, 0, 0)

            print(f"{question.answer}  {request.POST.get(f'q_{question.id}')}")
            if question.answer == request.POST.get(f'q_{question.id}'):
                correct += 1
            else:
                wrong += 1
                print(f"q_{question.id}")

        file.add_paragraph().add_run(
            f"Tog'ri javoblar: {correct}\n"
            f"Xato javoblar: {wrong}\n"
            f"Jami savollar: {total}\n"
            f"Vaqt: {total_seconds}"
        )

        user_first_name = request.user.username
        choose_exam_name = choose_exam.name.replace(' ', '_').replace('|', '')
        if temprary_first_name is not None:
            path = f"media/documents/{temprary_first_name}_{choose_exam_name}.docx"
            path_to = f"documents\{temprary_first_name}_{choose_exam_name}.docx"
        else:
            path = f"media/documents/{user_first_name}_{choose_exam_name}.docx"
            path_to = f"documents\{user_first_name}_{choose_exam_name}.docx"

        file.save(path)
        t = request.POST.get('timetest')
        w = request.POST.get("get_second")
        try:
            Results.objects.create(
                user=user,
                file=path_to,
                exam=choose_exam,
                correct=correct,
                wrong=wrong,
                time_out=total_seconds
            )

            Temporary_user.objects.create(
                first_name=user.first_name,
                last_name=user.last_name,
                file=path_to,
                exam_name=choose_exam,
                correct=correct,
                wrong=wrong,
                time_out=total_seconds
            )

        except:
            Temporary_user.objects.create(
                first_name=temprary_first_name,
                last_name=temprary_last_name,
                exam_name=choose_exam,
                file=path_to,
                correct=correct,
                wrong=wrong,
                time_out=total_seconds
            )

        context = {
            "total": total,
            "correct": correct,
            "wrong": wrong,
            "choose_exam": choose_exam,
            "path_to": path_to,
            "user": user,
            "duration": total_seconds,
        }
        return redirect('result_exam', pk=pk)

    else:
        context = {
            'total': [i for i in range(1, len(list(choose_exam_quizes))+1)],
            'choose_exam': choose_exam,
            'choose_exam_quizes': choose_exam_quizes,
            'time': choose_exam.get_all_time() * 60
        }
        return render(request, 'exams_items.html', context)


def result_exam(request, pk=None):
    choose_exam = models.Exams.objects.get(pk=pk)
    user = request.user
    try:
        results = Results.objects.filter(
            user=user,
            exam=choose_exam
        ).order_by('-id')[0]
    except:
        results = Temporary_user.objects.all().order_by('-id')[0]
        print(results)
        user = f"{results.first_name.title()} {results.last_name.title()}"
        print(user)
    context = {
        "wrong": results.wrong,
        "correct": results.correct,
        "duration": results.time_out,
        "total": results.wrong+results.correct,
        "path_to": results.file,
        "user": user
    }

    return render(request, 'results.html', context)


def ResultsListView(request, pk):
    choose_exam = models.Exams.objects.get(pk=pk)
    if request.method == "GET":
        return render(request, 'results_list.html')
    elif request.method == "POST":
        search_word = request.POST.get('q')
        if search_word is not None:
            search_word = str(search_word).strip()
            result = Q(first_name__contains=search_word) | \
                Q(last_name__contains=search_word)
            results_temrary_users = Temporary_user.objects.filter(
                exam_name=choose_exam).filter(result).order_by('-correct')

    else:
        results_temrary_users = Temporary_user.objects.filter(
            exam_name=choose_exam,
        ).order_by('-correct')

    context = {
        'choose_exam': choose_exam,
        'results_temrary_users': results_temrary_users
    }
    return render(request, 'results_list.html', context)


def download_page_view(request, path):
    file_path = os.path.join(settings.MEDIA_ROOT, path)
    if os.path.exists(file_path):
        with open(file_path, 'rb') as fh:
            response = HttpResponse(
                fh.read(), content_type="application/vnd.ms-word")
            response['Content-Disposition'] = 'inline; filename=' + \
                os.path.basename(file_path)
            return response
    raise Http404


def add_exam(request):
    text = ''
    if request.method == "POST":
        first_name = request.user.first_name
        last_name = request.user.last_name
        full_name = f"{first_name} {last_name}"
        exam_name = request.POST.get('exam_name')
        science_name = request.POST.get('science_name')
        time_limit = request.POST.get('time_limit')
        questions_count = request.POST.get('questions_count')
        status = request.POST.get('status')
        models.Exams.objects.create(
            author=full_name,
            name=exam_name,
            science_name=science_name,
            time_limit=time_limit,
            questions_count=questions_count,
            status=status
        )
        text = "Muaffaqiyatli imtihon yaratildi!"
        return redirect('profile')

    context = {
        'text': text,
    }
    return render(request, 'add_exam.html', context)


def add_quiz_view(request, pk):
    choose_exam = models.Exams.objects.get(pk=pk)
    text = ''
    if request.method == 'POST':
        user = request.user.get_full_name
        if request.POST.get('add_only'):
            question = request.POST.get('question')
            answer_a = request.POST.get('answer_a')
            answer_b = request.POST.get('answer_b')
            answer_c = request.POST.get('answer_c')
            answer_d = request.POST.get('answer_d')
            answer_id = request.POST.get('answer_id')
            print('*' * 20 + answer_id)
            models.Quiz.objects.create(
                exam=choose_exam,
                question=question,
                answer_a=answer_a,
                answer_b=answer_b,
                answer_c=answer_c,
                answer_d=answer_d,
                answer=answer_id
            ).save()
            text = 'Muaffaqiyatli test yaratildi!'
            return redirect('profile')

        elif request.POST.get('add_excel'):
            print(f"{'='*30} + {request.POST.get('add_excel')}")
            excel_file = request.FILES.get('excel_file')
            models.QuizExcel.objects.create(
                exam=choose_exam,
                file=excel_file
            ).save()

            temporary_choose_exam = f"{choose_exam}".replace(
                ' ', '_').replace('||', '')
            excel_book = openpyxl.open(
                f'./media/excelfiles/file_{temporary_choose_exam}/{excel_file}', read_only=True)
            sheet = excel_book.worksheets[0]
            print(sheet)
            for row in sheet.rows:
                if ((row[0].value != "Savollar") and (row[0].value is not None)) and (row[1].value is not None) and (row[2].value is not None) and (row[3].value is not None) and (row[4].value is not None) and (row[5].value is not None):
                    models.Quiz.objects.create(
                        exam=choose_exam,
                        question=row[0].value,
                        answer_a=row[1].value,
                        answer_b=row[2].value,
                        answer_c=row[3].value,
                        answer_d=row[4].value,
                        answer=row[5].value
                    ).save()
                    print('*'*10 + row[5].value)
            text = 'Muaffiqiyatli majmua qo\'shildi!'
            return redirect('profile')

    context = {
        'choose_exam': choose_exam,
        'answers': dict(models.ANSWERS),
        'text': text
    }

    return render(request, 'add_quiz.html', context)


def index_view(request):
    return render(request, 'index.html')


def exam_update(request, pk):
    exam = models.Exams.objects.get(pk=pk)
    if request.method == "POST":
        exam.exam_name = request.POST.get('exam_name')
        exam.science_name = request.POST.get('science_name')
        exam.time_limit = request.POST.get('time_limit')
        exam.questions_count = request.POST.get('questions_count')
        exam.status = request.POST.get('status')
        exam.save()
        return redirect("profile")
    context = {"exam": exam}
    return render(request, "exam_update.html", context)


def exam_delete(request, pk):
    exam = models.Exams.objects.get(pk=pk)
    if request.method == "POST":
        exam.delete()
        return redirect("profile")
    context = {"exam": exam}
    return render(request, "exam_delete.html", context)
